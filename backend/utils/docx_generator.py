from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

def generate_docx(data: dict, template_name: str) -> io.BytesIO:
    layout = data.get('template', 'modern')
    document = Document()
    
    personal = data.get('personal', {})
    
    # Fonts
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial' if layout != 'classic' else 'Times New Roman'
    font.size = Pt(11)

    # Header
    fullName = personal.get('fullName', 'Your Name')
    if layout == 'classic':
        h1 = document.add_paragraph()
        run = h1.add_run(fullName.upper())
        run.bold = True
        run.font.size = Pt(22)
        h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"{personal.get('jobTitle', '')}\n").italic = True
        p.add_run(f"{personal.get('email', '')} | {personal.get('phone', '')} | {personal.get('linkedin', '')}")
    else:
        h1 = document.add_heading(fullName, 0)
        h1.style.font.color.rgb = RGBColor(37, 99, 235) # Blue

        p = document.add_paragraph()
        p.add_run(f"{personal.get('jobTitle', '')}\n").bold = True
        p.add_run(f"{personal.get('email', '')} | {personal.get('phone', '')}")

    # Sections
    def add_section_header(text):
        h = document.add_heading(text, level=1)
        if layout == 'classic':
            h.style.font.color.rgb = RGBColor(0, 0, 0)
            h.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Add a line below
            p = document.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run("_" * 50)
            run.font.size = Pt(2)

    # Summary
    if personal.get('summary'):
        add_section_header('Profile Summary')
        document.add_paragraph(personal['summary'])

    # Experience
    if data.get('experience'):
        add_section_header('Professional Experience')
        for exp in data['experience']:
            p = document.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            title = p.add_run(exp['title'])
            title.bold = True
            
            p.add_run(f"\t{exp['date']}").italic = True
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT # Simple tab fallback
            
            p2 = document.add_paragraph()
            p2.add_run(exp['company']).italic = True
            
            document.add_paragraph(exp['description'])

    # Education
    if data.get('education'):
        add_section_header('Education')
        for edu in data['education']:
            p = document.add_paragraph()
            p.add_run(edu.get('degree', '')).bold = True
            p.add_run(f", {edu.get('school', '')} ({edu.get('date', '')})")

    # Skills
    if data.get('skills'):
        add_section_header('Skills')
        document.add_paragraph(', '.join(data['skills']))

    # Save to memory
    docx_stream = io.BytesIO()
    document.save(docx_stream)
    docx_stream.seek(0)
    return docx_stream
